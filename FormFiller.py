# Importing necessary libraries/modules 
import os  # Library for interacting with the operating system
import json  # Library for handling JSON data
import requests  # Library for making HTTP requests
import pandas as pd  # Library for data manipulation and analysis
from docx import Document  # Module for working with Word documents
from fuzzywuzzy import fuzz  # Module for fuzzy string matching
from openpyxl import load_workbook  # Module for working with Excel files
from pdf2docx import Converter  # Module for converting PDF to DOCX
import fitz  # Module for working with PDF documents
from collections import OrderedDict  # Ordered dictionary data structure
from docx2pdf import convert  # Module for converting DOCX to PDF

#These lines set up the necessary paths for input files (like forms and data), output directories, and API credentials. 

# Setting file paths and directories
form_file_path = '/Users/hope/Desktop/Final Project/24MARCH/Forms/Whitehorn_Scholarship_application_form_2023-24.docx'
data_file_path = '/Users/hope/Desktop/Final Project/Data/dummy_data.csv'
answers_file_path = '/Users/hope/Desktop/Final Project/saved_answers.json'
output_directory = '/Users/hope/Desktop/Final Project/OUTPUT_FINAL/'

# OpenAI API credentials
api_key = 'api-key-here'
selected_model = 'gpt-3.5-turbo'





#These functions handle saving and loading answers to/from a JSON file. The answers are stored as key-value pairs where the key is the question and the value is the answer.

# Function to save answers to a file
def save_answers_to_file(answers, file_path):
    with open(file_path, 'w') as file:
        json.dump(answers, file)

# Function to load answers from a file
def load_answers_from_file(file_path):
    try:
        with open(file_path, 'r') as file:
            # Use json.loads with object_pairs_hook to handle non-string keys
            return json.loads(file.read(), object_pairs_hook=OrderedDict)
    except FileNotFoundError:
        return None
    



#This function handles making requests to the ChatGPT API with retry logic to handle connection errors. It retries the request a maximum of 3 times before raising a ConnectionError.

# Function to call ChatGPT API with retry logic
def call_chat_gpt_with_retry(prompt, tokens, api_key, selected_model, max_retries=3):
    # API endpoint and headers
    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    # Payload for the API request
    payload = {
        "model": selected_model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7
    }

    retries = 0
    # Retry logic for handling connection errors 
    while retries < max_retries:
        try:
            response = requests.post(url, json=payload, headers=headers)
            return response.json()  # Ensures correct API response capture
        except requests.ConnectionError:
            retries += 1
            print(f"Retrying ({retries}/{max_retries}) due to ConnectionError...")
    
    raise ConnectionError("Max retries reached. Could not establish a connection.")



#This function cleans up and extracts the answer from the response obtained from the OpenAI API. It handles cases where the response may not contain the expected answer. 

# Function to clean up and extract the answer from OpenAI API response
def answer_cleanup(output):
    try:
        return output['choices'][0]['message']['content']
    except (KeyError, IndexError, TypeError):
        return "Error in processing the response."



#This function generates a persona description based on the context provided. It can be customized to generate different personas based on the context.

# Function to generate persona based on context
def generate_persona_for_form_filler(context):
    # Generates a persona based on the context provided
    if "medical" in context.lower():
        return "a meticulous form filler with expertise in healthcare and medical documentation"
    elif "finance" in context.lower():
        return "a detail-oriented form filler specializing in financial documentation"
    elif "university" in context.lower():
        return "a precise form filler with PHD level education applying for university."
    else:
        return "a meticulous form filler with expertise in general form filling and data entry"


#This function extracts the persona information based on the context. It uses the generate_persona_for_form_filler function to generate the persona description.

# Function to extract persona from the context
def extract_persona_info(context):
    return generate_persona_for_form_filler(context)




#This function generates answers for a list of questions using the provided context and saves the answers to a file. If saved answers are found, it uses those instead of generating new answers. 

# Function to generate answers for a list of questions using the provided context
def generate_answers_for_questions(context, questions, persona, api_key, selected_model, answers_file_path):
    saved_answers = load_answers_from_file(answers_file_path)

    if saved_answers:
        print("Using saved answers.")
        return saved_answers
    else:
        print("No saved answers found. Generating answers using GPT.")

    answers = {}
    persona = extract_persona_info(context)

    for question in questions:
        if hasattr(question, 'text'):
            question_text = question.text
        else:
            question_text = question

        prompt = f"Based on the following details: {context}. Persona: {persona}, what is the answer for '{question_text}'? Keep the response in mostly 1 word or few words, maximum of 5 words."

        response = call_chat_gpt_with_retry(prompt, 60, api_key, selected_model)
        cleaned_response = answer_cleanup(response)
        answers[question_text] = cleaned_response.strip()

    # Save the generated answers to a file
    save_answers_to_file(answers, answers_file_path)

    return answers



#This function fills in answers in a Word document based on a dictionary of keyword-answer pairs. It searches for keywords in the document and fills in the answers in the next paragraph or cell. 

# Function to fill in answers in a Word document based on a dictionary of keyword-answer pairs
def fill_in_answers_word(doc_path, answers_dict, output_path):
    # Open the Word document
    doc = Document(doc_path)
    updated = False  # Track if any changes were made

    # Process paragraphs in the document body
    for i, para in enumerate(doc.paragraphs):
        for keyword, answer in answers_dict.items():
            if keyword.lower() in para.text.lower():  # Case insensitive search for the keyword
                # Check if the next paragraph is empty and we are not at the last paragraph
                if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i + 1].text.strip():
                    # Insert the answer into the next paragraph
                    doc.paragraphs[i + 1].text = answer
                    updated = True

    # Process each table in the document
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for keyword, answer in answers_dict.items():
                    if keyword.lower() in cell.text.lower():
                        # Attempt to find the next empty cell in the row
                        if i + 1 < len(row.cells) and not row.cells[i + 1].text.strip():
                            row.cells[i + 1].text = answer
                            updated = True

    if updated:
        # Save the document to the specified output path
        doc.save(output_path)
        print("Document updated successfully.")
    else:
        print("No changes made to the document.")



#This function fills in answers in an Excel file based on a dictionary of keyword-answer pairs. It searches for keywords in the Excel file and fills in the answers in the next cell or row and saves the updated workbook to a new file.

# Function to fill in answers in an Excel file based on a dictionary of keyword-answer pairs
def fill_in_answers_excel_openpyxl(excel_path, answers_dict):
    # Load the original Excel workbook
    workbook = load_workbook(excel_path)

    # Process each sheet in the Excel file
    for sheet_name in workbook.sheetnames:
        # Get the sheet from the original workbook
        sheet = workbook[sheet_name]

        # Process each row in the sheet
        for row_num, row in enumerate(sheet.iter_rows(min_row=1), start=1):
            # Process each cell in the row
            for col_num, cell in enumerate(row, start=1):
                cell_value = cell.value
                for keyword, answer in answers_dict.items():
                    if keyword.lower() in str(cell_value).lower():
                        try:
                            # Check if there is a cell to the right (next to the keyword)
                            if col_num < sheet.max_column:
                                next_cell = sheet.cell(row=row_num, column=col_num + 1)
                                # Check if the cell is part of a merged range
                                if next_cell.coordinate in sheet.merged_cells:
                                    try:
                                        # Unmerge the cell to modify its value
                                        sheet.unmerge_cells(next_cell.coordinate)
                                    except KeyError:
                                        pass  # Ignore KeyError if the cell was not in merged_cells
                                # Set the value
                                next_cell.value = answer
                            # Check if there is a cell below (below the keyword)
                            elif row_num < sheet.max_row:
                                below_cell = sheet.cell(row=row_num + 1, column=col_num)
                                # Check if the cell is part of a merged range
                                if below_cell.coordinate in sheet.merged_cells:
                                    try:
                                        # Unmerge the cell to modify its value
                                        sheet.unmerge_cells(below_cell.coordinate)
                                    except KeyError:
                                        pass  # Ignore KeyError if the cell was not in merged_cells
                                # Set the value
                                below_cell.value = answer
                        except AttributeError as e:
                            print(f"Ignoring keyword '{keyword}' due to read-only cell: {e}")

    # Save the updated workbook
    output_file_path = os.path.join(output_directory, f"filled_{os.path.basename(excel_path)}")
    workbook.save(output_file_path)

    return output_file_path




#This function converts a PDF file to a DOCX file. It uses the pdf2docx library to perform the conversion.

# Function to convert PDF to DOCX
def convert_pdf_to_docx(pdf_file_path, output_docx_path):
    cv = Converter(pdf_file_path)
    cv.convert(output_docx_path, start=0, end=None)
    cv.close()



#This function converts a filled DOCX file back to PDF using the docx2pdf library. If the conversion is successful, it deletes the filled DOCX file. If there's an error during conversion, it prints an error message and removes the filled PDF output if it exists.

# Function to convert filled DOCX to PDF using docx2pdf library
def convert_docx_to_pdf(docx_file_path, output_pdf_path):
    try:
        # Convert filled DOCX back to PDF using docx2pdf library
        convert(docx_file_path, output_pdf_path)
        print("PDF conversion successful.")
        # Delete the filled Word document if the conversion is successful
        if os.path.exists(docx_file_path):
            os.remove(docx_file_path)
    except Exception as e:
        print(f"Error during PDF conversion: {e}")
        # If there's an error during conversion, remove the filled PDF output if it exists
        if os.path.exists(output_pdf_path):
            os.remove(output_pdf_path)



#This function reads text from a PDF file using the PyMuPDF library. It extracts text from each page of the PDF and concatenates it into a single string.

# Function to extract text from a PDF file using PyMuPDF
def read_pdf_file(file_path):
    doc = fitz.open(file_path)
    text = ''
    for page_num in range(doc.page_count):
        page = doc[page_num]
        text += page.get_text()
    return text




#This function extracts and filters questions from a PDF file. It reads the text from the PDF, extracts questions, filters duplicate questions, generates answers for the questions using the provided context, and fills in the answers in the PDF. 

# Function to extract and filter questions from PDF file
def extract_and_filter_questions_from_pdf(pdf_file_path, api_key, selected_model, data_context):
    answers_for_questions = {}

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1.text, question2.text) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question.text or '?' in question.text 

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    text_from_pdf = read_pdf_file(pdf_file_path)
    extracted_questions_from_pdf = extract_filtered_questions_from_text(text_from_pdf)
    filtered_questions = filter_duplicates(extracted_questions_from_pdf)

    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model)

    output_file_path = os.path.join(os.path.dirname(pdf_file_path), f"output_{os.path.basename(pdf_file_path)}")
    fill_in_answers_word(pdf_file_path, answers_for_questions, output_file_path)

    return filtered_questions




#This function extracts filtered questions from text based on the assumption that questions are marked with a colon, question mark, or period. It splits the text into lines and extracts lines that contain these punctuation marks.

# Function to extract filtered questions from text
def extract_filtered_questions_from_text(text):
    # Assuming questions are marked with a colon or question mark
    lines = text.split('\n')
    questions = []

    for line in lines:
        if ':' in line or '?' in line or '.' in line:
            questions.append(line.strip())

    return questions





#This function extracts and filters questions from a Word document. It defines helper functions for filtering duplicates and extracting questions from both DOCX and DOC formats. Then it generates answers for the questions and fills in the answers in the Word document.

# Function to extract and filter questions from Word document
def extract_and_filter_questions_from_word(form_file_path, api_key, selected_model):
    answers_for_questions = {}

    # Define helper functions for filtering duplicates and extracting questions

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1.text, question2.text) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question.text or '?' in question.text 

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    def extract_filtered_questions_from_docx(doc):
        extracted_questions = []

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        current_question = paragraph

                        if contains_colon_or_question_mark(current_question):
                            extracted_questions.append(current_question)

        return extracted_questions

    def extract_filtered_questions_from_doc(doc):
        extracted_questions = []

        for paragraph in doc.paragraphs:
            if contains_colon_or_question_mark(paragraph):
                extracted_questions.append(paragraph)

        return extracted_questions

    doc = Document(form_file_path)

    extracted_questions_from_docx = extract_filtered_questions_from_docx(doc)
    extracted_questions_from_doc = extract_filtered_questions_from_doc(doc)

    filtered_questions = filter_duplicates(extracted_questions_from_docx + extracted_questions_from_doc)

    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model, answers_file_path)

    output_file_path = os.path.join(os.path.dirname(output_directory), f"filled_{os.path.basename(form_file_path)}")
    fill_in_answers_word(form_file_path, answers_for_questions, output_file_path)

    return filtered_questions, answers_for_questions







#This function extracts and filters questions from an Excel file. It defines helper functions for filtering duplicates and extracting questions. Then it reads the Excel file, flattens the DataFrame to a list of strings (questions), filters duplicate questions, generates answers for the questions, and fills in the answers in the Excel file. 

# Function to extract and filter questions from Excel file
def extract_and_filter_questions_from_excel(excel_file_path, api_key, selected_model):
    answers_for_questions = {}

    # Define helper functions for filtering duplicates and extracting questions

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1, question2) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question or '?' in question

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    # Read Excel file using pandas
    df = pd.read_excel(excel_file_path)

    # Flatten the DataFrame to a list of strings (questions)
    questions_from_excel = [str(cell) for col in df.columns for cell in df[col]]

    filtered_questions = filter_duplicates(questions_from_excel)

    # Generate answers using GPT-3 API
    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model, answers_file_path)


    fill_in_answers_excel_openpyxl(excel_file_path, answers_for_questions)

    return filtered_questions




#This function reads input data from a file based on the file extension. It supports reading text from TXT, CSV, Word (DOCX, DOC), PDF, and Excel (XLSX, XLS) files.

# Function to read input data from txt, csv, or word
def read_data_file(file_path):
    _, file_extension = os.path.splitext(file_path)

    if file_extension.lower() == '.txt':
        with open(file_path, 'r') as file:
            return file.read()
    elif file_extension.lower() == '.csv':
        df = pd.read_csv(file_path)
        data_dict = df.to_dict(orient='records')[0]
        return '. '.join([f'"{key}": {value}' for key, value in data_dict.items()])
    elif file_extension.lower() == '.pdf':
        return read_pdf_file(file_path)
    elif file_extension.lower() in ['.docx', '.doc']:
        return read_word_file(file_path)
    elif file_extension.lower() == '.xlsx':
        return read_excel_file(file_path)
    else:
        print("Unsupported format, only txt, csv, docx, doc, pdf, and xlsx files.")
        return None




#This function reads text from a Word file using the python-docx library. It reads the text from each paragraph in the Word document and concatenates it into a single string.

# Function to read text from Word file
def read_word_file(file_path):
    doc = Document(file_path)
    return '\n'.join(paragraph.text for paragraph in doc.paragraphs)


#This function reads text from an Excel file using the pandas library. It reads the Excel file into a DataFrame and converts the first row of the DataFrame to a dictionary.

# Function to read text from Excel file
def read_excel_file(file_path):
    df = pd.read_excel(file_path)
    return df.to_dict(orient='records')[0]



#This function reads text from a PDF file using the PyMuPDF library. It extracts text from each page of the PDF and concatenates it into a single string.

# Function to read text from PDF file
def read_pdf_file(file_path):
    doc = fitz.open(file_path)
    text = ''
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text += page.get_text()
        return text



#This function cleans up a key by removing colons and extra spaces. 

def clean_key(key):
    return key.replace(':', '').strip()

#This function extracts text from a paragraph and removes leading and trailing whitespaces.

def extract_text(paragraph):
    return paragraph.text.strip()

#This function counts the number of filled answers and ignored questions based on the answers and questions provided. It cleans the keys of the answers and compares them with the cleaned keys of the questions to determine if an answer is filled or a question is ignored. 

def count_filled_and_ignored(answers, questions):
    filled_count = 0
    ignored_count = 0

    cleaned_answers = {clean_key(key): value for key, value in answers.items()}

    for question in questions:
        if isinstance(question, str):
            question_text = question
        else:
            question_text = extract_text(question)

        question_key = clean_key(question_text)

        if question_key in cleaned_answers:
            filled_count += 1
        else:
            ignored_count += 1
            print(f"Ignored Question: {question_key}")

    return filled_count, ignored_count



#This part of the code demonstrates the usage of the functions defined earlier. It reads the input data, extracts and filters questions based on the file format, generates answers for the questions, fills in the answers, and counts the filled and ignored questions. Finally, it prints the extracted questions, generated answers, and counts of filled and ignored questions.

data_context = read_data_file(data_file_path)

if data_context and form_file_path.lower().endswith(('.docx', '.doc')):
    questions, answers_for_questions = extract_and_filter_questions_from_word(form_file_path, api_key, selected_model)
    temp_docx_path = form_file_path

    output_file_path_filled = os.path.join(output_directory, f"filled_{os.path.basename(form_file_path)}")
    output_file_path_copy = os.path.join(output_directory, f"copy_{os.path.basename(form_file_path)}")

    fill_in_answers_word(temp_docx_path, answers_for_questions, output_file_path_filled)

    os.system(f'cp "{form_file_path}" "{output_file_path_copy}"')

elif form_file_path.lower().endswith('.pdf'):
    # Copy the original PDF file
    original_pdf_copy_path = os.path.join(output_directory, f"copy_{os.path.basename(form_file_path)}")
    os.system(f'cp "{form_file_path}" "{original_pdf_copy_path}"')

    # Convert PDF to DOCX
    temp_docx_path = os.path.join(output_directory, f"temp_{os.path.splitext(os.path.basename(form_file_path))[0]}.docx")
    convert_pdf_to_docx(form_file_path, temp_docx_path)

    # Extract and fill answers in the DOCX
    questions, answers_for_questions = extract_and_filter_questions_from_word(temp_docx_path, api_key, selected_model)

    # Output path for the filled DOCX
    output_file_path_filled = os.path.join(output_directory, f"filled_{os.path.basename(form_file_path)}")

    # Fill in answers in the DOCX
    fill_in_answers_word(temp_docx_path, answers_for_questions, output_file_path_filled)

    # Convert filled DOCX back to PDF
    filled_pdf_output_path = os.path.join(output_directory, f"filled_{os.path.splitext(os.path.basename(form_file_path))[0]}.pdf")
    convert_docx_to_pdf(output_file_path_filled, filled_pdf_output_path)

    # Delete the temporary DOCX file if it exists
    if temp_docx_path and os.path.exists(temp_docx_path):
        os.remove(temp_docx_path)

elif data_context and form_file_path.lower().endswith(('.xlsx', '.xls')):
    questions = extract_and_filter_questions_from_excel(form_file_path, api_key, selected_model)
    temp_docx_path = None

    answers_for_questions = generate_answers_for_questions(data_context, questions, "", api_key, selected_model, answers_file_path)

    output_file_path_excel = os.path.join(output_directory, f"filled_{os.path.basename(form_file_path)}")

    fill_in_answers_excel_openpyxl(form_file_path, answers_for_questions)

    os.system(f'cp "{form_file_path}" "{os.path.join(output_directory, "_copy_" + os.path.basename(form_file_path))}"')

else:
    print("Invalid form file format. Supported formats: Word Files (docx, doc), Excel Files (xls,xlsx), and PDF Files.")
    questions, answers_for_questions = None, None

if questions:
    print("Extracted Questions:")
    for question in questions:
        if hasattr(question, 'text'):
            print(question.text)
        else:
            print(question)

    print("\nGenerated Answers:")
    for question, answer in answers_for_questions.items():
        print(f"{question}: {answer}\n")



# Count the number of filled answers and ignored questions
        
filled_count, ignored_count = count_filled_and_ignored(answers_for_questions, questions)

print(f"Filled Answers: {filled_count}")
print(f"Ignored Questions: {ignored_count}")