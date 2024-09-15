from docx import Document
from fuzzywuzzy import fuzz
import os
from textblob import TextBlob

def extract_questions_from_word(word_file_path):
    def is_duplicate(question1, question2):
        return fuzz.ratio(question1.text, question2.text) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question.text or '?' in question.text

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    def extract_questions_from_docx(doc):
        extracted_questions = []

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        current_question = paragraph

                        if contains_colon_or_question_mark(current_question):
                            extracted_questions.append(current_question)

        for shape in doc.inline_shapes:
            if hasattr(shape, 'text_frame') and shape.text_frame:
                text = shape.text_frame.text

                if contains_colon_or_question_mark(text):
                    extracted_questions.append(shape)

        return extracted_questions

    def extract_questions_from_doc(doc):
        extracted_questions = []

        for paragraph in doc.paragraphs:
            if contains_colon_or_question_mark(paragraph):
                extracted_questions.append(paragraph)

        return extracted_questions

    _, file_extension = os.path.splitext(word_file_path)

    if file_extension.lower() == '.docx':
        doc = Document(word_file_path)
        extracted_questions_from_docx = extract_questions_from_docx(doc)
    elif file_extension.lower() == '.doc':
        print("This file has to be converted and uploaded again. Use doc_to_docx.py to convert.")
        return []
    else:
        print("Unsupported format, only docx and doc files.")
        return []

    extracted_questions_from_doc = extract_questions_from_doc(doc)

    filtered_questions = filter_duplicates(extracted_questions_from_docx + extracted_questions_from_doc)

    return filtered_questions

# Example usage
word_file_path = '/Users/hope/Desktop/Final Project/Data/BCU-application-form.docx'
questions = extract_questions_from_word(word_file_path)

if questions:
    print("Extracted Questions:")
    for question in questions:
        print(question.text)

def get_sentiment(sentence):
    analysis = TextBlob(sentence)
    return analysis.sentiment.polarity

def generate_answers(questions):
    answers = []

    for question in questions:
        sentiment = get_sentiment(question.text)
        if sentiment > 0:
            answers.append("Yes")
        else:
            answers.append("No")

    return answers


if questions:
    print("Sentiment analysis Extracted Questions:")
    for question in questions:
        print(question.text)

    answers = generate_answers(questions)

    print("\nGenerated Answers:")
    for question, answer in zip(questions, answers):
        print(f"{question.text}: {answer}")
 