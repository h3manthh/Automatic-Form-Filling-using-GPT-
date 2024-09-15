from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
import os
from docx import Document
from fuzzywuzzy import fuzz
import requests
from retrying import retry
import pandas as pd
from pdf2docx import Converter
import fitz
from openpyxl import load_workbook
import openpyxl

app = Flask(__name__)

data_context = None

# Function to clear saved answers for a keyword
def clear_saved_answers(keyword):
    file_path = f'{keyword}_answers.json'
    
    try:
        os.remove(file_path)
        print(f"Cleared saved answers for keyword '{keyword}'.")
    except FileNotFoundError:
        print(f"No saved answers found for keyword '{keyword}'.")

# Example usage:
clear_saved_answers("temp_keyword")

# Function to save answers to a keyword
def save_answers_to_keyword(keyword, answers):
    # with open(f'{keyword}_answers.json', 'w') as file:
    #     json.dump(answers, file)
    pass

# Function to load answers from a keyword
def load_answers_from_keyword(keyword):
    # try:
    #     with open(f'{keyword}_answers.json', 'r') as file:
    #         return json.load(file)
    # except FileNotFoundError:
    return None

def call_chat_gpt_with_retry(prompt, tokens, api_key, selected_model, keyword, max_retries=3):
    saved_answers = load_answers_from_keyword(keyword)
    if saved_answers:
        print(f"Using saved answers for keyword '{keyword}': {saved_answers}")
        return {'choices': [{'message': {'content': answer}} for answer in saved_answers]}

    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    payload = {
        "model": selected_model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7
    }

    retries = 0
    while retries < max_retries:
        try:
            response = requests.post(url, json=payload, headers=headers)
            result = response.json()
            answers = [result['choices'][0]['message']['content']]
            #save_answers_to_keyword(keyword, answers)
            return result
        except requests.ConnectionError:
            retries += 1
            print(f"Retrying ({retries}/{max_retries}) due to ConnectionError...")

    raise ConnectionError("Max retries reached. Could not establish a connection.")

# Function to clean up and extract the answer from OpenAI API response
def answer_cleanup(output):
    try:
        return output['choices'][0]['message']['content']
    except (KeyError, IndexError, TypeError):
        return "Error in processing the response."
    
# Function to generate persona based on context
#def generate_persona_for_form_filler(context):
    if "medical" in context.lower():
        return "a meticulous form filler with expertise in healthcare and medical documentation"
    elif "finance" in context.lower():
        return "a detail-oriented form filler specializing in financial documentation"
    elif "university" in context.lower():
        return "a precise form filler with PHD level education applying for university."
    else:
        return "a meticulous form filler with expertise in general form filling and data entry"

# Function to extract persona from the context
def extract_persona_info(context):
    if context:
        if "medical" in context.lower():
            return "a meticulous form filler with expertise in healthcare and medical documentation"
        elif "finance" in context.lower():
            return "a detail-oriented form filler specializing in financial documentation"
        elif "university" in context.lower():
            return "a precise form filler with PHD level education applying for university."
        else:
            return "a meticulous form filler with expertise in general form filling and data entry"
    else:
        return "a form filler with no specific context"

# Function to generate answers for a list of questions using the provided context
def generate_answers_for_questions(context, questions, persona, api_key, selected_model):
    answers = {}
    persona = extract_persona_info(context)

    for question in questions:
        if hasattr(question, 'text'):
            # If question is an object with 'text' attribute (e.g., from Word), use question.text
            prompt = f"Based on the following details: {context}. Persona: {persona}, what is the answer for '{question.text}'? Keep the response in mostly 1 word or few words, maximum of 5 words."
        else:
            # If question is a string (e.g., from Excel), use it directly
            prompt = f"Based on the following details: {context}. Persona: {persona}, what is the answer for '{question}'? Keep the response in mostly 1 word or few words, maximum of 5 words."

        response = call_chat_gpt_with_retry(prompt, 60, api_key, selected_model, "temp_keyword") # Example token count; adjust as needed
        cleaned_response = answer_cleanup(response)
        answers[question] = cleaned_response.strip()

    return answers

# Function to extract and filter questions from PDF
def extract_and_filter_questions_from_pdf(pdf_file_path, api_key, selected_model):
    answers_for_questions = {}

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1.text, question2.text) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question.text or '?' in question.text

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    def read_pdf_file(file_path):
        doc = fitz.open(file_path)
        text = ''
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text += page.get_text()
        return text

    text_from_pdf = read_pdf_file(pdf_file_path)
    extracted_questions_from_pdf = extract_filtered_questions_from_text(text_from_pdf)
    filtered_questions = filter_duplicates(extracted_questions_from_pdf)

    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model)

    output_file_path = f"./output/{secure_filename(pdf_file_path)}_filled.docx"
    fill_in_answers(pdf_file_path, answers_for_questions, output_file_path)

    return filtered_questions

def extract_filtered_questions_from_text(text):
    lines = text.split('\n')
    questions = []

    for line in lines:
        if ':' in line or '?' in line:
            questions.append(line.strip())

    return questions

# Function to extract and filter questions from Word document
def extract_and_filter_questions_from_word(word_file_path, api_key, selected_model):
    answers_for_questions = {}

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1.text, question2.text) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question.text or '?' in question.text

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    #def replace_with_answer(paragraph, answer):
        text = paragraph.text
        if ':' in text:
            question_text = text.split(':', 1)[0]
            paragraph.text = f"{question_text}: {answer}"

    def extract_filtered_questions_from_docx(doc):
        extracted_questions = []

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        current_question = paragraph

                        if contains_colon_or_question_mark(current_question):
                            extracted_questions.append(current_question)

        for shape in doc.inline_shapes:
            if shape.type == 3: # Check if it's a text box or not
                if hasattr(shape, 'text_frame'):
                    text = shape.text_frame.text

                    if contains_colon_or_question_mark(text):
                        extracted_questions.append(shape)
            else:
                pass

        return extracted_questions

    def extract_filtered_questions_from_doc(doc):
        extracted_questions = []

        for paragraph in doc.paragraphs:
            if contains_colon_or_question_mark(paragraph):
                extracted_questions.append(paragraph)

        return extracted_questions

    doc = Document(word_file_path)

    extracted_questions_from_docx = extract_filtered_questions_from_docx(doc)
    extracted_questions_from_doc = extract_filtered_questions_from_doc(doc)

    filtered_questions = filter_duplicates(extracted_questions_from_docx + extracted_questions_from_doc)

    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model)
    #for question in filtered_questions:
        #replace_with_answer(question, answers_for_questions.get(question.text, ''))

    output_file_path = f"./output/{secure_filename(word_file_path)}_filled.docx"
    doc.save(output_file_path)

    return filtered_questions

# Function to extract and filter questions from Excel file
def extract_and_filter_questions_from_excel(excel_file_path, api_key, selected_model):
    answers_for_questions = {}

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1, question2) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question or '?' in question

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    # Read Excel file using pandas
    df = pd.read_excel(excel_file_path)

    # Flatten the DataFrame to a list of strings (questions)
    questions_from_excel = [str(cell) for col in df.columns for cell in df[col]]

    filtered_questions = filter_duplicates(questions_from_excel)

    # Generate answers using GPT-3 API
    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model)

    # Specify the path for the output Excel file after filling in answers
    output_file_path_excel = f"./output/{secure_filename(excel_file_path)}_filled.xlsx"

    # Create a copy of the original Excel DataFrame
    df_filled = df.copy()

    # Replace the corresponding cells with answers
    for question, answer in answers_for_questions.items():
        df_filled.replace(question, answer, inplace=True)

    # Convert 'Not applicable.' to NaN in the DataFrame
    df_filled.replace('Not applicable.', float('nan'), inplace=True)    

    # Save the filled DataFrame to a new Excel file
    df_filled.to_excel(output_file_path_excel, index=False)

    return filtered_questions

# Function to read input data from txt, csv, or word
def read_input_data(file_path):
    _, file_extension = os.path.splitext(file_path)

    if file_extension.lower() == '.txt':
        return read_text_file(file_path)
    elif file_extension.lower() == '.csv':
        df = pd.read_csv(file_path)
        data_dict = df.to_dict(orient='records')[0]
        return '. '.join([f'"{key}": {value}' for key, value in data_dict.items()])
    elif file_extension.lower() == '.pdf':
        temp_docx_path = f"./temp/{secure_filename(file_path)}.docx"
        convert_pdf_to_docx(file_path, temp_docx_path)
        return read_word_file(temp_docx_path)
    elif file_extension.lower() in ['.docx', '.doc']:
        return read_word_file(file_path)
    elif file_extension.lower() == '.xlsx':
        return read_excel_file(file_path)
    else:
        print("Unsupported format, only txt, csv, docx, doc, pdf, and xlsx files.")
        return None

def read_text_file(file_path):
    with open(file_path, 'r') as file:
        return file.read()

def read_word_file(file_path):
    doc = Document(file_path)
    return '\n'.join(paragraph.text for paragraph in doc.paragraphs)

def read_excel_file(file_path):
    df = pd.read_excel(file_path)
    return df.to_dict(orient='records')[0]

# Function to fill in answers in a Word document based on a dictionary of keyword-answer pairs
def fill_in_answers(doc_path, answers_dict, output_path):
    doc = Document(doc_path)

    # Process paragraphs in the document body
    for i, para in enumerate(doc.paragraphs):
        for keyword, answer in answers_dict.items():
            if isinstance(para, Document):  # Check if it's a Paragraph object
                if keyword.lower() in para.text.lower():  # Case insensitive search for the keyword
                    # Check if the next paragraph is empty and we are not at the last paragraph
                    if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i + 1].text.strip():
                        # Insert the answer into the next paragraph
                        doc.paragraphs[i + 1].text = answer

    # Process each table in the document
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for keyword, answer in answers_dict.items():
                    if isinstance(cell, Document):  # Check if it's a Paragraph object
                        if keyword.lower() in cell.text.lower():
                            # Attempt to find the next empty cell in the row
                            if i + 1 < len(row.cells) and not row.cells[i + 1].text.strip():
                                row.cells[i + 1].text = answer

    # Save the document to the specified output path
    doc.save(output_path)

# Function to fill in answers in an Excel file based on a dictionary of keyword-answer pairs
def fill_in_answers_excel_openpyxl(excel_path, answers_dict):
    # Load the original Excel workbook
    workbook = openpyxl.load_workbook(excel_path)

    # Process each sheet in the Excel file
    for sheet_name in workbook.sheetnames:
        # Get the sheet from the original workbook
        sheet = workbook[sheet_name]

        # Process each row in the sheet
        for row_num, row in enumerate(sheet.iter_rows(min_row=1), start=1):
            # Process each cell in the row
            for col_num, cell in enumerate(row, start=1):
                cell_value = cell.value
                for keyword, answer in answers_dict.items():
                    if keyword.lower() in str(cell_value).lower():
                        try:
                            # Check if there is a cell to the right (next to the keyword)
                            if col_num < sheet.max_column:
                                next_cell = sheet.cell(row=row_num, column=col_num + 1)
                                # Check if the cell is part of a merged range
                                if next_cell.coordinate in sheet.merged_cells:
                                    try:
                                        # Unmerge the cell to modify its value
                                        sheet.unmerge_cells(next_cell.coordinate)
                                    except KeyError:
                                        pass  # Ignore KeyError if the cell was not in merged_cells
                                # Set the value
                                next_cell.value = answer
                            # Check if there is a cell below (below the keyword)
                            elif row_num < sheet.max_row:
                                below_cell = sheet.cell(row=row_num + 1, column=col_num)
                                # Check if the cell is part of a merged range
                                if below_cell.coordinate in sheet.merged_cells:
                                    try:
                                        # Unmerge the cell to modify its value
                                        sheet.unmerge_cells(below_cell.coordinate)
                                    except KeyError:
                                        pass  # Ignore KeyError if the cell was not in merged_cells
                                # Set the value
                                below_cell.value = answer
                        except AttributeError as e:
                            print(f"Ignoring keyword '{keyword}' due to read-only cell: {e}")

    # Save the updated workbook
    output_path = f"./output/{secure_filename(excel_path)}_filled.xlsx"
    workbook.save(output_path)

    return output_path

# Function to convert PDF to DOCX
def convert_pdf_to_docx(pdf_file_path, output_docx_path):
    cv = Converter(pdf_file_path)
    cv.convert(output_docx_path, start=0, end=None)
    cv.close()

# Function to process form submission
def process_form_submission(input_file, data_file):
    api_key = 'api key'  # Replace with your actual OpenAI API key
    selected_model = 'gpt-3.5-turbo'

    global data_context  # Declare data_context as a global variable
    
    data_context = read_input_data(data_file)

    if data_context and input_file.lower().endswith(('.docx', '.doc')):
        questions = extract_and_filter_questions_from_word(input_file, api_key, selected_model)
    elif data_context and input_file.lower().endswith('.pdf'):
        temp_docx_path = f"./temp/{secure_filename(input_file)}.docx"
        convert_pdf_to_docx(input_file, temp_docx_path)
        questions = extract_and_filter_questions_from_word(temp_docx_path, api_key, selected_model)
    elif data_context and input_file.lower().endswith(('.xlsx', '.xls')):
        questions = extract_and_filter_questions_from_excel(input_file, api_key, selected_model)
    else:
        return "Invalid file format. Supported formats: txt, csv, docx, doc, pdf, xls, and xlsx."

    answers_for_questions = generate_answers_for_questions(data_context, questions, "", api_key, selected_model)

    output_file_path = f"./output/{secure_filename(input_file)}_filled.docx"
    fill_in_answers(input_file, answers_for_questions, output_file_path)

    return output_file_path

# Route for the home page
@app.route('/', methods=['GET', 'POST'])
def home():
    if request.method == 'POST':
        if 'input_file' not in request.files or 'data_file' not in request.files:
            return render_template('index.html', error="Please select both input and data files.")

        input_file = request.files['input_file']
        data_file = request.files['data_file']

        if input_file.filename == '' or data_file.filename == '':
            return render_template('index.html', error="Please select both input and data files.")

        input_file_path = f"./uploads/{secure_filename(input_file.filename)}"
        data_file_path = f"./uploads/{secure_filename(data_file.filename)}"
        input_file.save(input_file_path)
        data_file.save(data_file_path)

        output_file_path = process_form_submission(input_file_path, data_file_path)

        return send_file(output_file_path, as_attachment=True)

    return render_template('index.html')

if __name__ == '__main__':
    # Create the 'uploads' directory if it doesn't exist
    os.makedirs('uploads', exist_ok=True)
    
    # Create the 'temp' directory if it doesn't exist
    os.makedirs('temp', exist_ok=True)
    
    # Create the 'output' directory if it doesn't exist
    os.makedirs('output', exist_ok=True)
    app.run(debug=True)