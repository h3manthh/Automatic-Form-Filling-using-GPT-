from docx import Document
from fuzzywuzzy import fuzz
import openai
import os

# Set your OpenAI GPT-3.5 API key
openai.api_key = 'api key'

def extract_questions_from_word(word_file_path):
    def contains_colon_or_question_mark(text):
        return ':' in text or '?' in text

    def extract_questions_from_docx(doc):
        extracted_questions = []

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if contains_colon_or_question_mark(text):
                            extracted_questions.append(paragraph)

        for shape in doc.inline_shapes:
            if hasattr(shape, 'text_frame') and shape.text_frame:
                text = shape.text_frame.text
                if contains_colon_or_question_mark(text):
                    extracted_questions.append(shape)

        return extracted_questions

    def extract_questions_from_doc(doc):
        extracted_questions = []

        for paragraph in doc.paragraphs:
            text = paragraph.text
            if contains_colon_or_question_mark(text):
                extracted_questions.append(paragraph)

        return extracted_questions

    _, file_extension = os.path.splitext(word_file_path)

    if file_extension.lower() == '.docx':
        doc = Document(word_file_path)
        extracted_questions_from_docx = extract_questions_from_docx(doc)
    elif file_extension.lower() == '.doc':
        print("This file has to be converted and uploaded again. Use doc_to_docx.py to convert.")
        return []
    else:
        print("Unsupported format, only docx and doc files.")
        return []

    extracted_questions_from_doc = extract_questions_from_doc(doc)

    return extracted_questions_from_docx + extracted_questions_from_doc

def generate_answer_based_on_prompt(prompt):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=150  # Adjust the max tokens as needed
    )
    answer = response.choices[0].message['content'].strip()
    return answer

# Example usage
word_file_path = '/Users/hope/Desktop/Final Project/PGT_Readmission_form_Dec_2020.docx'
questions = extract_questions_from_word(word_file_path)

if questions:
    print("Extracted Questions:")
    for question in questions:
        print(question.text)

    # Example prompt
    prompt = "Assuming this is to fill a form, print out the sentence that would seem like a question when filling a form.\n\n"
    for question in questions:
        prompt += f"- {question.text}\n"
        answer = generate_answer_based_on_prompt(prompt)

    print("\nGenerated Answer:")
    print(answer)
    
