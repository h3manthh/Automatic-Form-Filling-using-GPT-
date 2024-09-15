from docx import Document

def fill_tables_with_text(file_path, text_file_path):
    try:
        doc = Document(file_path)
        word_list = []

        # Read words from the text file
        with open(text_file_path, 'r') as text_file:
            word_list = text_file.read().split()

        print(f"Words from text file: {word_list}")

        # Get the number of tables in the document
        num_tables = len(doc.tables)
        print(f"Number of tables in the document: {num_tables}")

        # Counter for filled tables
        filled_tables_count = 0

        # Iterate over tables in the document
        for table in doc.tables:
            # Iterate over rows and cells in the table
            for row in table.rows:
                for cell in row.cells:
                    if word_list:
                        # Pop a word from the list and fill the cell
                        cell.text = word_list.pop(0)

            # Increment the counter for each table filled
            filled_tables_count += 1

        doc.save(file_path)
        print(f"Tables filled successfully in: {file_path}")
        print(f"Number of tables filled: {filled_tables_count}")

    except Exception as e:
        print(f"An error occurred: {e}")

# Example Usage:
file_path = "/Users/hope/Desktop/Final Project/Data/BCU-application-form.docx"
text_file_path = "/Users/hope/Desktop/Final Project/Data/random text.txt"

fill_tables_with_text(file_path, text_file_path)
