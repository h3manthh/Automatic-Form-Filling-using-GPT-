from docx import Document
from fuzzywuzzy import fuzz
import pandas as pd
from pdf2docx import Converter 
import os
from openai import OpenAI

class Analyzer:
    def __init__(self) -> None:

        pass

    def extract(self, file_path):
        raise NotImplementedError("Method 'extract' must be implemented in subclass")

    def analyze(self, file_path):
        raise NotImplementedError("Method 'analyze' must be implemented in subclass")


class ExcelAnalyzer(Analyzer):
    def __init__(self) -> None:
        super().__init__()
    
    def is_duplicate(self, question1, question2):
        return fuzz.ratio(question1, question2) > 80  # Adjust the threshold as needed
    
    def filter_duplicates(self, questions):
        filtered_questions = []
        for current_question in questions:
            if not any(self.is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    
    def extract(self, file_path):
        df = pd.read_excel(file_path)
        questions_from_excel = [str(cell) for col in df.columns for cell in df[col]]
        filtered_questions = self.filter_duplicates(questions_from_excel)

        questions = filtered_questions
        return questions

    def analyze(self, file_path):
        questions = self.extract(file_path)
        return questions
        pass

class WordAnalyzer(Analyzer):
    def __init__(self) -> None:
        super().__init__()

    def is_duplicate(self, question1, question2):
        return fuzz.ratio(question1.text, question2.text) > 800  # Adjust the threshold as needed
    
    def contains_colon_or_question_mark(self, question):
        return ':' in question.text or '?' in question.text or '.' in question.text
    
    def extract_filtered_questions_from_tables(self, doc):
        extracted_questions = []

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        current_question = paragraph

                        if self.contains_colon_or_question_mark(current_question):
                            extracted_questions.append(current_question)

        return extracted_questions
    
    def extract_filtered_questions_from_paragraphs(self, doc):
        extracted_questions = []

        for paragraph in doc.paragraphs:
            if self.contains_colon_or_question_mark(paragraph):
                extracted_questions.append(paragraph)

        return extracted_questions
    
    def filter_duplicates(self, questions):
        filtered_questions = []
        for current_question in questions:
            if not any(self.is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    
    def extract(self, file_path) -> None:
        doc = Document(file_path)

        extracted_questions_from_tables = self.extract_filtered_questions_from_tables(doc)
        extracted_questions_from_paragraphs = self.extract_filtered_questions_from_paragraphs(doc)

        filtered_questions = self.filter_duplicates(extracted_questions_from_tables + extracted_questions_from_paragraphs)
        questions = [question.text if hasattr(question, 'text') else question for question in filtered_questions]

        return questions

    def analyze(self, file_path) -> None:
        questions = self.extract(file_path)

        return questions
    

class PdfAnalyzer(WordAnalyzer):
    def __init__(self) -> None:
        super().__init__()
    
    def convert_pdf_to_docx(self, pdf_file_path, output_docx_path):
        cv = Converter(pdf_file_path)
        cv.convert(output_docx_path, start=0, end=None)
        cv.close()
    
    def extract(self, file_path):
        temp_docx_path = os.path.join(os.path.dirname(file_path), "temp_output_file.docx")
        self.convert_pdf_to_docx(file_path, temp_docx_path)

        return super().extract(temp_docx_path)
    
    def analyze(self, file_path):
        return super().analyze(file_path)