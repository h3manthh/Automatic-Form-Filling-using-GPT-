# app.py
from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
from flask import send_from_directory
import os
from docx import Document
from fuzzywuzzy import fuzz
import requests
from retrying import retry
import pandas as pd
from pdf2docx import Converter
import fitz
import json
from openpyxl import load_workbook
from openpyxl.styles import NamedStyle
from openpyxl import Workbook
import xlsxwriter
import openpyxl
from docx.text.paragraph import Paragraph
from collections import OrderedDict

answers_for_questions = {}


app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'txt', 'pdf', 'docx', 'doc', 'xlsx', 'xls', 'csv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def save_answers_to_file(answers, file_path):
    with open(file_path, 'w') as file:
        json.dump(answers, file)

def load_answers_from_file(file_path):
    try:
        with open(file_path, 'r') as file:
            return json.loads(file.read(), object_pairs_hook=OrderedDict)
    except FileNotFoundError:
        return None

def clean_key(key):
    # Remove colons and extra spaces
    return key.replace(':', '').strip()
def extract_text(paragraph):
    # Extract text from paragraph object
    return paragraph.text.strip()

def count_filled_and_ignored(answers, questions):
    filled_count = 0
    ignored_count = 0

    cleaned_answers = {clean_key(key): value for key, value in answers.items()}

    for question in questions:
        # Extract text from paragraph object
        question_text = extract_text(question)

        # Convert the question text to a string before using it as a key
        question_key = clean_key(question_text)

        if question_key in cleaned_answers:
            filled_count += 1
        else:
            ignored_count += 1
            print(f"Ignored Question: {question_key}")

    

    return filled_count, ignored_count

# Use the original questions without conversion to string
filled_count, ignored_count = count_filled_and_ignored(answers_for_questions, questions)

print(f"Filled Answers: {filled_count}")
print(f"Ignored Questions: {ignored_count}")

def call_chat_gpt_with_retry(prompt, tokens, api_key, selected_model, max_retries=3):
    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    payload = {
        "model": selected_model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7
    }

    retries = 0
    while retries < max_retries:
        try:
            response = requests.post(url, json=payload, headers=headers)
            return response.json()  # Ensures correct API response capture
        except requests.ConnectionError:
            retries += 1
            print(f"Retrying ({retries}/{max_retries}) due to ConnectionError...")
    
    raise ConnectionError("Max retries reached. Could not establish a connection.")

# Function to clean up and extract the answer from OpenAI API response
def answer_cleanup(output):
    try:
        return output['choices'][0]['message']['content']
    except (KeyError, IndexError, TypeError):
        return "Error in processing the response."

# Function to generate persona based on context
def generate_persona_for_form_filler(context):
    if "medical" in context.lower():
        return "a meticulous form filler with expertise in healthcare and medical documentation"
    elif "finance" in context.lower():
        return "a detail-oriented form filler specializing in financial documentation"
    elif "university" in context.lower():
        return "a precise form filler with PHD level education applying for university."
    else:
        return "a meticulous form filler with expertise in general form filling and data entry"

# Function to extract persona from the context
def extract_persona_info(context):
    return generate_persona_for_form_filler(context)

# Function to generate answers for a list of questions using the provided context
def generate_answers_for_questions(context, questions, persona, api_key, selected_model, answers_file_path):
    saved_answers = load_answers_from_file(answers_file_path)

    if saved_answers:
        print("Using saved answers.")
        return saved_answers
    else:
        print("No saved answers found. Generating answers using GPT.")

    answers = {}
    persona = extract_persona_info(context)

    for question in questions:
        if hasattr(question, 'text'):
            # If question is an object with 'text' attribute (e.g., from Word), use question.text
            prompt = f"Based on the following details: {context}. Persona: {persona}, what is the answer for '{question.text}'? Keep the response in mostly 1 word or few words, maximum of 5 words."
        else:
            # If question is a string (e.g., from Excel), use it directly
            prompt = f"Based on the following details: {context}. Persona: {persona}, what is the answer for '{question}'? Keep the response in mostly 1 word or few words, maximum of 5 words."

        response = call_chat_gpt_with_retry(prompt, 60, api_key, selected_model)  # Example token count; adjust as needed
        cleaned_response = answer_cleanup(response)
        answers[question.text] = cleaned_response.strip()
        
    # Save the generated answers to a file
    save_answers_to_file(answers, answers_file_path)

    return answers

# Function to fill in answers in a Word document based on a dictionary of keyword-answer pairs
def fill_in_answers_word(doc_path, answers_dict, output_path):
    doc = Document(doc_path)
    updated = False  # Track if any changes were made

    # Process paragraphs in the document body
    for i, para in enumerate(doc.paragraphs):
        for keyword, answer in answers_dict.items():
            if keyword.lower() in para.text.lower():  # Case insensitive search for the keyword
                # Check if the next paragraph is empty and we are not at the last paragraph
                if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i + 1].text.strip():
                    # Insert the answer into the next paragraph
                    doc.paragraphs[i + 1].text = answer
                    updated = True

    # Process each table in the document
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for keyword, answer in answers_dict.items():
                    if keyword.lower() in cell.text.lower():
                        # Attempt to find the next empty cell in the row
                        if i + 1 < len(row.cells) and not row.cells[i + 1].text.strip():
                            row.cells[i + 1].text = answer
                            updated = True

    if updated:
        # Save the document to the specified output path
        doc.save(output_path)
        print("Document updated successfully.")
    else:
        print("No changes made to the document.")

# Function to fill in answers in an Excel file based on a dictionary of keyword-answer pairs
def fill_in_answers_excel_openpyxl(excel_path, answers_dict):
    # Load the original Excel workbook
    workbook = openpyxl.load_workbook(excel_path)

    # Process each sheet in the Excel file
    for sheet_name in workbook.sheetnames:
        # Get the sheet from the original workbook
        sheet = workbook[sheet_name]

        # Process each row in the sheet
        for row_num, row in enumerate(sheet.iter_rows(min_row=1), start=1):
            # Process each cell in the row
            for col_num, cell in enumerate(row, start=1):
                cell_value = cell.value
                for keyword, answer in answers_dict.items():
                    if keyword.lower() in str(cell_value).lower():
                        try:
                            # Check if there is a cell to the right (next to the keyword)
                            if col_num < sheet.max_column:
                                next_cell = sheet.cell(row=row_num, column=col_num + 1)
                                # Check if the cell is part of a merged range
                                if next_cell.coordinate in sheet.merged_cells:
                                    try:
                                        # Unmerge the cell to modify its value
                                        sheet.unmerge_cells(next_cell.coordinate)
                                    except KeyError:
                                        pass  # Ignore KeyError if the cell was not in merged_cells
                                # Set the value
                                next_cell.value = answer
                            # Check if there is a cell below (below the keyword)
                            elif row_num < sheet.max_row:
                                below_cell = sheet.cell(row=row_num + 1, column=col_num)
                                # Check if the cell is part of a merged range
                                if below_cell.coordinate in sheet.merged_cells:
                                    try:
                                        # Unmerge the cell to modify its value
                                        sheet.unmerge_cells(below_cell.coordinate)
                                    except KeyError:
                                        pass  # Ignore KeyError if the cell was not in merged_cells
                                # Set the value
                                below_cell.value = answer
                        except AttributeError as e:
                            print(f"Ignoring keyword '{keyword}' due to read-only cell: {e}")

    # Save the updated workbook
    output_path = '/Users/hope/Desktop/Final Project/FormFiller/output/Mar5_temp_2_file_filled.xlsx'
    workbook.save(output_path)

    return output_path

# Function to convert PDF to DOCX
def convert_pdf_to_docx(pdf_file_path, output_docx_path):
    cv = Converter(pdf_file_path)
    cv.convert(output_docx_path, start=0, end=None)
    cv.close()

# Function to extract text from a PDF file using PyMuPDF
def read_pdf_file(file_path):
    doc = fitz.open(file_path)
    text = ''
    for page_num in range(doc.page_count):
        page = doc[page_num]
        text += page.get_text()
    return text

def extract_and_filter_questions_from_pdf(pdf_file_path, api_key, selected_model):
    answers_for_questions = {}

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1.text, question2.text) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question.text or '?' in question.text

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    text_from_pdf = read_pdf_file(pdf_file_path)
    extracted_questions_from_pdf = extract_filtered_questions_from_text(text_from_pdf)
    filtered_questions = filter_duplicates(extracted_questions_from_pdf)

    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model)

    output_file_path = '/Users/hope/Desktop/Final Project/Output/Final Output/feb28_3.docx'
    fill_in_answers_word(pdf_file_path, answers_for_questions, output_file_path)

    return filtered_questions
def extract_filtered_questions_from_text(text):
    # Assuming questions are marked with a colon or question mark
    lines = text.split('\n')
    questions = []

    for line in lines:
        if ':' in line or '?' in line:
            questions.append(line.strip())

    return questions

def extract_and_filter_questions_from_word(word_file_path, api_key, selected_model):
    answers_for_questions = {}

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1.text, question2.text) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question.text or '?' in question.text

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    #def replace_with_answer(paragraph, answer):
        text = paragraph.text
        if ':' in text:
            question_text = text.split(':', 1)[0]
            paragraph.text = f"{question_text}: {answer}"

    def extract_filtered_questions_from_docx(doc):
        extracted_questions = []

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        current_question = paragraph

                        if contains_colon_or_question_mark(current_question):
                            extracted_questions.append(current_question)

        for shape in doc.inline_shapes:
            if shape.type == 3:  # Check if it's a text box or not
                if hasattr(shape, 'text_frame'):
                    text = shape.text_frame.text

                    if contains_colon_or_question_mark(text):
                        extracted_questions.append(shape)
            else:
            # Process other types of inline shapes if needed
                pass

        return extracted_questions


    def extract_filtered_questions_from_doc(doc):
        extracted_questions = []

        for paragraph in doc.paragraphs:
            if contains_colon_or_question_mark(paragraph):
                extracted_questions.append(paragraph)

        return extracted_questions

    doc = Document(word_file_path)

    extracted_questions_from_docx = extract_filtered_questions_from_docx(doc)
    extracted_questions_from_doc = extract_filtered_questions_from_doc(doc)

    filtered_questions = filter_duplicates(extracted_questions_from_docx + extracted_questions_from_doc)

    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model, answers_file_path)
   # for question in filtered_questions:
       # replace_with_answer(question, answers_for_questions.get(question.text, ''))

    output_file_path = '/Users/hope/Desktop/Final Project/FormFiller/output/feb28_3.docx'
    doc.save(output_file_path)

    return filtered_questions, answers_for_questions  # Add this line to return the list of filtered questions

# Function to extract and filter questions from Excel file
def extract_and_filter_questions_from_excel(excel_file_path, api_key, selected_model):
    answers_for_questions = {}

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1, question2) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question or '?' in question

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    # Read Excel file using pandas
    df = pd.read_excel(excel_file_path)

    # Flatten the DataFrame to a list of strings (questions)
    questions_from_excel = [str(cell) for col in df.columns for cell in df[col]]

    filtered_questions = filter_duplicates(questions_from_excel)

    # Generate answers using GPT-3 API
    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model, answers_file_path)


    # Specify the path for the output Excel file after filling in answers
    output_file_path_excel = '/Users/hope/Desktop/Final Project/Output/March Output/Mar5_file_filled.xlsx'

    # Create a copy of the original Excel DataFrame
    df_filled = df.copy()

    # Replace the corresponding cells with answers
    for question, answer in answers_for_questions.items():
        df_filled.replace(question, answer, inplace=True)

    # Convert 'Not applicable.' to NaN in the DataFrame
    df_filled.replace('Not applicable.', float('nan'), inplace=True)    

    # Save the filled DataFrame to a new Excel file
    df_filled.to_excel(output_file_path_excel, index=False)

    return filtered_questions

# Function to read input data from txt, csv, or word
def read_input_data(file_path):
    _, file_extension = os.path.splitext(file_path)

    if file_extension.lower() == '.txt':
        return read_text_file(file_path)
    elif file_extension.lower() == '.csv':
        df = pd.read_csv(file_path)
        data_dict = df.to_dict(orient='records')[0]
        return '. '.join([f'"{key}": {value}' for key, value in data_dict.items()])
    elif file_extension.lower() == '.pdf':
        # Convert PDF to DOCX
        temp_docx_path = '/Users/hope/Desktop/Final Project/FormFiller/output/Mar2_temp_output_file.docx'  # Provide a temporary path for the converted DOCX
        convert_pdf_to_docx(file_path, temp_docx_path)
        return read_word_file(temp_docx_path)
    elif file_extension.lower() in ['.docx', '.doc']:
        return read_word_file(file_path)
    elif file_extension.lower() == '.xlsx':
        return read_excel_file(file_path)
    else:
        print("Unsupported format, only txt, csv, docx, doc, pdf, and xlsx files.")
        return None
    
def read_text_file(file_path):
    with open(file_path, 'r') as file:
        return file.read()

def read_word_file(file_path):
    doc = Document(file_path)
    return '\n'.join(paragraph.text for paragraph in doc.paragraphs)

# Function to read Excel file and return text
def read_excel_file(file_path):
    df = pd.read_excel(file_path)
    return df.to_dict(orient='records')[0]


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    # Check if the post request has the file part
    if 'data_file' not in request.files or 'input_file' not in request.files:
        return render_template('result.html', result='Error: Files not provided.')

    data_file = request.files['data_file']
    input_file = request.files['input_file']

    # Check if files are empty
    if data_file.filename == '' or input_file.filename == '':
        return render_template('result.html', result='Error: Files not provided.')

    # Check if files have allowed extensions
    if not (allowed_file(data_file.filename) and allowed_file(input_file.filename)):
        return render_template('result.html', result='Error: Invalid file extension.')

    # Save uploaded files
    data_file_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(data_file.filename))
    input_file_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(input_file.filename))
    data_file.save(data_file_path)
    input_file.save(input_file_path)

    # Execute your processing code here (use functions from your existing code)

    # Return a response or redirect to a result page
    result = 'Processing complete!'
    filled_count, ignored_count = count_filled_and_ignored(answers_for_questions, questions)
    result += f"\nFilled Answers: {filled_count}\nIgnored Questions: {ignored_count}"

    return render_template('result.html', result=result)

@app.route('/Users/hope/Desktop/Final Project/FormFiller/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

if __name__ == '__main__':
    app.run(debug=True)