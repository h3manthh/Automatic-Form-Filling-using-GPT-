from flask import Flask, render_template, request, send_file
import os
import json
import requests
import pandas as pd
from docx import Document
from fuzzywuzzy import fuzz
from openpyxl import load_workbook
from openpyxl.styles import NamedStyle
from pdf2docx import Converter 
import fitz
from collections import OrderedDict

app = Flask(__name__)

# Function to save answers to a file
def save_answers_to_file(answers, file_path):
    with open(file_path, 'w') as file:
        json.dump(answers, file)

# Function to load answers from a file
def load_answers_from_file(file_path):
    try:
        with open(file_path, 'r') as file:
            # Use json.loads with object_pairs_hook to handle non-string keys
            return json.loads(file.read(), object_pairs_hook=OrderedDict)
    except FileNotFoundError:
        return None
    
# Function to call ChatGPT API with retry logic
def call_chat_gpt_with_retry(prompt, tokens, api_key, selected_model, max_retries=3):
    url = "https://api.openai.com/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    payload = {
        "model": selected_model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7
    }

    retries = 0
    while retries < max_retries:
        try:
            response = requests.post(url, json=payload, headers=headers)
            return response.json()  # Ensures correct API response capture
        except requests.ConnectionError:
            retries += 1
            print(f"Retrying ({retries}/{max_retries}) due to ConnectionError...")
    
    raise ConnectionError("Max retries reached. Could not establish a connection.")

# Function to clean up and extract the answer from OpenAI API response
def answer_cleanup(output):
    try:
        return output['choices'][0]['message']['content']
    except (KeyError, IndexError, TypeError):
        return "Error in processing the response."

# Function to generate persona based on context
def generate_persona_for_form_filler(context):
    if "medical" in context.lower():
        return "a meticulous form filler with expertise in healthcare and medical documentation"
    elif "finance" in context.lower():
        return "a detail-oriented form filler specializing in financial documentation"
    elif "university" in context.lower():
        return "a precise form filler with PHD level education applying for university."
    else:
        return "a meticulous form filler with expertise in general form filling and data entry"

# Function to extract persona from the context
def extract_persona_info(context):
    return generate_persona_for_form_filler(context)

# Function to generate answers for a list of questions using the provided context
def generate_answers_for_questions(context, questions, persona, api_key, selected_model, answers_file_path):
    saved_answers = load_answers_from_file(answers_file_path)

    if saved_answers:
        print("Using saved answers.")
        return saved_answers
    else:
        print("No saved answers found. Generating answers using GPT.")

    answers = {}
    persona = extract_persona_info(context)

    for question in questions:
        if hasattr(question, 'text'):
            # If question is an object with 'text' attribute (e.g., from Word), use question.text
            question_text = question.text
        else:
            # If question is a string (e.g., from Excel), use it directly
            question_text = question

        if hasattr(question, 'text'):
            prompt = f"Based on the following details: {context}. Persona: {persona}, Give me the answer for, '{question_text}' as follows: STRICT RULES: 1. Do not answer for the questions that are, general descriptive text or narrative explanations or any content that does not directly request user input or response. 2. Answer for the remaining questions as follows, answers mostly as a consise answer, preferably a singly word or a short answer. For creative questions or for questions there is no data try to create a story from the data provided make it compelling for the question."
        else:
            prompt = f"Based on the following details: {context}. Persona: {persona}, what is the answer for '{question_text}'? Keep the response in mostly 1 word or few words, maximum of 5 words."

        response = call_chat_gpt_with_retry(prompt, 60, api_key, selected_model)  # Example token count; adjust as needed
        cleaned_response = answer_cleanup(response)
        answers[question_text] = cleaned_response.strip()
        
    # Save the generated answers to a file
    save_answers_to_file(answers, answers_file_path)

    return answers

# Function to fill in answers in a Word document based on a dictionary of keyword-answer pairs
def fill_in_answers_word(doc_path, answers_dict, output_path):
    doc = Document(doc_path)
    updated = False  # Track if any changes were made

    # Process paragraphs in the document body
    for i, para in enumerate(doc.paragraphs):
        for keyword, answer in answers_dict.items():
            if keyword.lower() in para.text.lower():  # Case insensitive search for the keyword
                # Check if the next paragraph is empty and we are not at the last paragraph
                if i + 1 < len(doc.paragraphs) and not doc.paragraphs[i + 1].text.strip():
                    # Insert the answer into the next paragraph
                    doc.paragraphs[i + 1].text = answer
                    updated = True

    # Process each table in the document
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                for keyword, answer in answers_dict.items():
                    if keyword.lower() in cell.text.lower():
                        # Attempt to find the next empty cell in the row
                        if i + 1 < len(row.cells) and not row.cells[i + 1].text.strip():
                            row.cells[i + 1].text = answer
                            updated = True

    if updated:
        # Save the document to the specified output path
        doc.save(output_path)
        print("Document updated successfully.")
    else:
        print("No changes made to the document.")

# Function to fill in answers in an Excel file based on a dictionary of keyword-answer pairs
def fill_in_answers_excel_openpyxl(excel_path, answers_dict):
    # Load the original Excel workbook
    workbook = load_workbook(excel_path)

    # Process each sheet in the Excel file
    for sheet_name in workbook.sheetnames:
        # Get the sheet from the original workbook
        sheet = workbook[sheet_name]

        # Process each row in the sheet
        for row_num, row in enumerate(sheet.iter_rows(min_row=1), start=1):
            # Process each cell in the row
            for col_num, cell in enumerate(row, start=1):
                cell_value = cell.value
                for keyword, answer in answers_dict.items():
                    if keyword.lower() in str(cell_value).lower():
                        try:
                            # Check if there is a cell to the right (next to the keyword)
                            if col_num < sheet.max_column:
                                next_cell = sheet.cell(row=row_num, column=col_num + 1)
                                # Check if the cell is part of a merged range
                                if next_cell.coordinate in sheet.merged_cells:
                                    try:
                                        # Unmerge the cell to modify its value
                                        sheet.unmerge_cells(next_cell.coordinate)
                                    except KeyError:
                                        pass  # Ignore KeyError if the cell was not in merged_cells
                                # Set the value
                                next_cell.value = answer
                            # Check if there is a cell below (below the keyword)
                            elif row_num < sheet.max_row:
                                below_cell = sheet.cell(row=row_num + 1, column=col_num)
                                # Check if the cell is part of a merged range
                                if below_cell.coordinate in sheet.merged_cells:
                                    try:
                                        # Unmerge the cell to modify its value
                                        sheet.unmerge_cells(below_cell.coordinate)
                                    except KeyError:
                                        pass  # Ignore KeyError if the cell was not in merged_cells
                                # Set the value
                                below_cell.value = answer
                        except AttributeError as e:
                            print(f"Ignoring keyword '{keyword}' due to read-only cell: {e}")

    # Save the updated workbook
    output_file_path = os.path.join(output_directory, f"filled_{os.path.basename(excel_path)}")
    workbook.save(output_file_path)

    return output_file_path

# Function to convert PDF to DOCX
def convert_pdf_to_docx(pdf_file_path, output_docx_path):
    cv = Converter(pdf_file_path)
    cv.convert(output_docx_path, start=0, end=None)
    cv.close()

# Function to extract text from a PDF file using PyMuPDF
def read_pdf_file(file_path):
    doc = fitz.open(file_path)
    text = ''
    for page_num in range(doc.page_count):
        page = doc[page_num]
        text += page.get_text()
    return text

# Function to extract and filter questions from PDF file
def extract_and_filter_questions_from_pdf(pdf_file_path, api_key, selected_model):
    answers_for_questions = {}

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1.text, question2.text) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question.text or '?' in question.text or '.' in question.text

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    text_from_pdf = read_pdf_file(pdf_file_path)
    extracted_questions_from_pdf = extract_filtered_questions_from_text(text_from_pdf)
    filtered_questions = filter_duplicates(extracted_questions_from_pdf)

    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model)

    output_file_path = os.path.join(os.path.dirname(pdf_file_path), f"output_{os.path.basename(pdf_file_path)}")
    fill_in_answers_word(pdf_file_path, answers_for_questions, output_file_path)

    return filtered_questions

# Function to extract filtered questions from text
def extract_filtered_questions_from_text(text):
    # Assuming questions are marked with a colon or question mark
    lines = text.split('\n')
    questions = []

    for line in lines:
        if ':' in line or '?' in line or '.' in line:
            questions.append(line.strip())

    return questions

# Function to extract and filter questions from Word document
def extract_and_filter_questions_from_word(word_file_path, api_key, selected_model):
    answers_for_questions = {}

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1.text, question2.text) > 800  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question.text or '?' in question.text or '.' in question.text

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    def extract_filtered_questions_from_docx(doc):
        extracted_questions = []

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        current_question = paragraph

                        if contains_colon_or_question_mark(current_question):
                            extracted_questions.append(current_question)

        return extracted_questions

    def extract_filtered_questions_from_doc(doc):
        extracted_questions = []

        for paragraph in doc.paragraphs:
            if contains_colon_or_question_mark(paragraph):
                extracted_questions.append(paragraph)

        return extracted_questions

    doc = Document(word_file_path)

    extracted_questions_from_docx = extract_filtered_questions_from_docx(doc)
    extracted_questions_from_doc = extract_filtered_questions_from_doc(doc)

    filtered_questions = filter_duplicates(extracted_questions_from_docx + extracted_questions_from_doc)

    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model, answers_file_path)

    output_file_path = os.path.join(os.path.dirname(word_file_path), f"output_{os.path.basename(word_file_path)}")
    fill_in_answers_word(word_file_path, answers_for_questions, output_file_path)

    return filtered_questions, answers_for_questions

# Function to extract and filter questions from Excel file
def extract_and_filter_questions_from_excel(excel_file_path, api_key, selected_model):
    answers_for_questions = {}

    def is_duplicate(question1, question2):
        return fuzz.ratio(question1, question2) > 80  # Adjust the threshold as needed

    def contains_colon_or_question_mark(question):
        return ':' in question or '?' in question

    def filter_duplicates(questions):
        filtered_questions = []
        for current_question in questions:
            if not any(is_duplicate(current_question, q) for q in filtered_questions):
                filtered_questions.append(current_question)
        return filtered_questions

    # Read Excel file using pandas
    df = pd.read_excel(excel_file_path)

    # Flatten the DataFrame to a list of strings (questions)
    questions_from_excel = [str(cell) for col in df.columns for cell in df[col]]

    filtered_questions = filter_duplicates(questions_from_excel)

    # Generate answers using GPT-3 API
    answers_for_questions = generate_answers_for_questions(data_context, filtered_questions, "", api_key, selected_model, answers_file_path)


    fill_in_answers_excel_openpyxl(excel_file_path, answers_for_questions)

    return filtered_questions

# Function to read input data from txt, csv, or word file
def read_input_data_file(file_path):
    if file_path.endswith('.txt'):
        with open(file_path, 'r') as file:
            data_context = file.read()
    elif file_path.endswith('.csv'):
        df = pd.read_csv(file_path)
        data_context = '\n'.join([str(cell) for col in df.columns for cell in df[col]])
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        data_context = '\n'.join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format. Only .txt, .csv, and .docx are supported.")

    return data_context

# Function to download a file
def download_file(file_path):
    return send_file(file_path, as_attachment=True)

# Function to upload a file
def upload_file(request):
    if 'file' not in request.files:
        return None
    file = request.files['file']
    if file.filename == '':
        return None
    return file

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    if 'file' not in request.files:
        return render_template('index.html', error='No file uploaded.')

    file = request.files['file']
    if file.filename == '':
        return render_template('index.html', error='No file uploaded.')

    if file:
        file.save(file.filename)
        if file.filename.endswith('.pdf'):
            filtered_questions = extract_and_filter_questions_from_pdf(file.filename)
        elif file.filename.endswith('.docx'):
            filtered_questions, answers_for_questions = extract_and_filter_questions_from_word(file.filename)
        elif file.filename.endswith('.xlsx'):
            filtered_questions = extract_and_filter_questions_from_excel(file.filename)
        else:
            return render_template('index.html', error='Unsupported file format. Only PDF, DOCX, and XLSX are supported.')

        return render_template('result.html', questions=filtered_questions)

    return render_template('index.html', error='Error processing file.')

if __name__ == '__main__':
    app.run(debug=True)
